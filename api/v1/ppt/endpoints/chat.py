import base64
from io import BytesIO
import json
import os
import re
import tempfile
from typing import List, Literal, Optional
import uuid

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from google import genai
from sqlalchemy.ext.asyncio import AsyncSession

from services.documents_loader import DocumentsLoader
from services.credit_service import consume_credits, get_request_user_id
from services.database import get_async_session
from utils.get_env import get_google_api_key_env, get_google_model_env
from utils.asset_directory_utils import get_images_directory


CHAT_ROUTER = APIRouter(prefix="/chat", tags=["Chat"])


class ChatHistoryItem(BaseModel):
    role: Literal["user", "assistant"]
    content: str


class ChatMessageResponse(BaseModel):
    success: bool
    reply: str
    used_documents: List[str] = []
    used_images: int = 0
    remaining_credits: Optional[int] = None


class ChatImageResponse(BaseModel):
    success: bool
    image_url: str
    remaining_credits: Optional[int] = None


EssayFormat = Literal["APA", "MLA", "CUSTOM"]


def _unique_models(models: List[Optional[str]]) -> List[str]:
    unique: List[str] = []
    for model in models:
        if model and model not in unique:
            unique.append(model)
    return unique


def _normalize_model_name(model: Optional[str]) -> Optional[str]:
    if not model:
        return None
    normalized = model.strip()
    if normalized.startswith("models/"):
        return normalized.split("/", 1)[1]
    return normalized


def _parse_history(history_raw: Optional[str]) -> List[ChatHistoryItem]:
    if not history_raw:
        return []
    try:
        parsed = json.loads(history_raw)
        if not isinstance(parsed, list):
            return []
        validated: List[ChatHistoryItem] = []
        for item in parsed:
            validated.append(ChatHistoryItem(**item))
        return validated
    except Exception:
        return []


@CHAT_ROUTER.post("/message", response_model=ChatMessageResponse)
async def chat_message(
    request: Request,
    message: str = Form(...),
    model: Optional[str] = Form(None),
    history: Optional[str] = Form(None),
    document_files: Optional[List[UploadFile]] = File(None),
    image_files: Optional[List[UploadFile]] = File(None),
    sql_session: AsyncSession = Depends(get_async_session),
):
    api_key = get_google_api_key_env() or os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GOOGLE_API_KEY is not configured")

    if not message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")
    user_id = get_request_user_id(request)

    try:
        remaining_credits = await consume_credits(sql_session, user_id, "chat")
        client = genai.Client(api_key=api_key)
        history_items = _parse_history(history)

        contents = []
        for item in history_items[-8:]:
            contents.append(
                {
                    "role": "model" if item.role == "assistant" else "user",
                    "parts": [{"text": item.content}],
                }
            )

        doc_context_chunks: List[str] = []
        used_documents: List[str] = []
        if document_files:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_paths: List[str] = []
                for doc in document_files:
                    if not doc.filename:
                        continue
                    doc_path = os.path.join(temp_dir, doc.filename)
                    with open(doc_path, "wb") as f:
                        f.write(await doc.read())
                    temp_paths.append(doc_path)
                    used_documents.append(doc.filename)

                if temp_paths:
                    loader = DocumentsLoader(file_paths=temp_paths)
                    await loader.load_documents(temp_dir=temp_dir, load_text=True, load_images=False)
                    for idx, text in enumerate(loader.documents):
                        if not text:
                            continue
                        filename = used_documents[idx] if idx < len(used_documents) else f"Document {idx + 1}"
                        doc_context_chunks.append(f"[{filename}]\n{text[:4000]}")

        user_parts = []
        doc_context = "\n\n".join(doc_context_chunks).strip()
        if doc_context:
            user_parts.append(
                {
                    "text": (
                        "Use the following document context to answer accurately.\n\n"
                        f"{doc_context}\n\n"
                        f"User message: {message}"
                    )
                }
            )
        else:
            user_parts.append({"text": message})

        used_images = 0
        if image_files:
            for image_file in image_files:
                raw = await image_file.read()
                if not raw:
                    continue
                mime_type = image_file.content_type or "image/png"
                user_parts.append(
                    {
                        "inline_data": {
                            "mime_type": mime_type,
                            "data": base64.b64encode(raw).decode("utf-8"),
                        }
                    }
                )
                used_images += 1

        contents.append({"role": "user", "parts": user_parts})

        chat_models = _unique_models(
            [
                _normalize_model_name(model),
                _normalize_model_name(get_google_model_env()),
                "gemini-2.5-flash",
                "gemini-1.5-flash",
            ]
        )
        response = None
        last_error: Optional[Exception] = None
        for model_name in chat_models:
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=contents,
                    config={"response_mime_type": "text/plain"},
                )
                if getattr(response, "text", None):
                    break
            except Exception as model_error:
                last_error = model_error
                continue

        if response is None:
            error_msg = str(last_error) if last_error else "No model response received"
            raise HTTPException(status_code=502, detail=f"Chat model failed: {error_msg}")

        reply = (getattr(response, "text", None) or "").strip()
        if not reply:
            raise HTTPException(status_code=502, detail="Model returned an empty response")

        return ChatMessageResponse(
            success=True,
            reply=reply,
            used_documents=used_documents,
            used_images=used_images,
            remaining_credits=remaining_credits,
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat request failed: {str(e)}")


@CHAT_ROUTER.get("/generate-image", response_model=ChatImageResponse)
async def generate_chat_image(
    request: Request,
    prompt: str,
    model: Optional[str] = None,
    sql_session: AsyncSession = Depends(get_async_session),
):
    api_key = get_google_api_key_env() or os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GOOGLE_API_KEY is not configured")
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt cannot be empty")
    user_id = get_request_user_id(request)

    try:
        remaining_credits = await consume_credits(sql_session, user_id, "image")
        client = genai.Client(api_key=api_key)
        image_path: Optional[str] = None
        image_models = _unique_models(
            [
                _normalize_model_name(model),
                "gemini-2.5-flash-image-preview",
                "gemini-3-pro-image-preview",
            ]
        )
        last_error: Optional[Exception] = None
        for model_name in image_models:
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=[{"role": "user", "parts": [{"text": prompt}]}],
                )
            except Exception as model_error:
                last_error = model_error
                continue

            for candidate in response.candidates or []:
                content = getattr(candidate, "content", None)
                if not content:
                    continue
                for part in content.parts or []:
                    if getattr(part, "inline_data", None) is not None:
                        image = part.as_image()
                        filename = f"{uuid.uuid4()}.jpg"
                        image_path = os.path.join(get_images_directory(), filename)
                        image.save(image_path)
                        break
                if image_path:
                    break
            if image_path:
                break

        if not image_path:
            error_msg = str(last_error) if last_error else "Gemini did not return an image"
            raise HTTPException(status_code=502, detail=f"Image generation failed: {error_msg}")

        return ChatImageResponse(
            success=True,
            image_url=f"/app_data/images/{os.path.basename(image_path)}",
            remaining_credits=remaining_credits,
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Image generation failed: {str(e)}")


def _essay_prompt(
    prompt: str,
    essay_format: EssayFormat,
    school_name: Optional[str],
    course_code: Optional[str],
    course_name: Optional[str],
    student_name: Optional[str],
    instructor_name: Optional[str],
    due_date: Optional[str],
    custom_instructions: Optional[str],
) -> str:
    header_context = (
        f"School: {school_name or '[blank]'}\n"
        f"Course Code: {course_code or '[blank]'}\n"
        f"Course Name: {course_name or '[blank]'}\n"
        f"Student Name: {student_name or '[blank]'}\n"
        f"Instructor: {instructor_name or '[blank]'}\n"
        f"Due Date: {due_date or '[blank]'}"
    )
    return (
        f"Write a complete academic essay in {essay_format} format.\n"
        "Output rules:\n"
        "1) First line must be the essay title only.\n"
        "2) Then one blank line.\n"
        "3) Then full essay body in clear paragraphs.\n"
        "4) Include in-text citations where suitable and include a final references/works-cited section.\n"
        "5) Do not use markdown.\n\n"
        f"Metadata for cover/header:\n{header_context}\n\n"
        f"Essay request:\n{prompt}\n\n"
        f"Additional requirements:\n{custom_instructions or 'None'}"
    )


def _split_title_and_body(text: str) -> tuple[str, List[str]]:
    lines = [line.rstrip() for line in text.splitlines()]
    non_empty = [line for line in lines if line.strip()]
    if not non_empty:
        return ("Essay", [""])
    title = non_empty[0].strip()[:180] or "Essay"
    body_text = text[text.find(non_empty[0]) + len(non_empty[0]) :].strip()
    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [body_text] if body_text else [""]
    return (title, paragraphs)


def _split_body_and_references(paragraphs: List[str]) -> tuple[List[str], List[str]]:
    if not paragraphs:
        return [""], []

    ref_pattern = re.compile(r"^(references|works\s*cited|bibliography)\s*$", re.IGNORECASE)
    ref_idx: Optional[int] = None

    for idx, para in enumerate(paragraphs):
        first_line = para.strip().splitlines()[0].strip() if para.strip() else ""
        if ref_pattern.match(first_line):
            ref_idx = idx
            break

    if ref_idx is None:
        return paragraphs, []

    body_paragraphs = paragraphs[:ref_idx]
    references_paragraphs = paragraphs[ref_idx:]

    if not body_paragraphs:
        body_paragraphs = [""]

    cleaned_references: List[str] = []
    for idx, para in enumerate(references_paragraphs):
        if idx == 0:
            lines = [line for line in para.splitlines() if line.strip()]
            if len(lines) <= 1:
                continue
            cleaned_references.append("\n".join(lines[1:]).strip())
            continue
        cleaned_references.append(para)

    cleaned_references = [p for p in cleaned_references if p.strip()]
    return body_paragraphs, cleaned_references


@CHAT_ROUTER.post("/essay")
async def generate_essay_docx(
    request: Request,
    prompt: str = Form(...),
    essay_format: EssayFormat = Form("APA"),
    school_name: Optional[str] = Form(None),
    course_code: Optional[str] = Form(None),
    course_name: Optional[str] = Form(None),
    student_name: Optional[str] = Form(None),
    instructor_name: Optional[str] = Form(None),
    due_date: Optional[str] = Form(None),
    custom_instructions: Optional[str] = Form(None),
    template_docx: Optional[UploadFile] = File(None),
    sql_session: AsyncSession = Depends(get_async_session),
):
    api_key = get_google_api_key_env() or os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GOOGLE_API_KEY is not configured")
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Essay prompt cannot be empty")
    user_id = get_request_user_id(request)

    try:
        await consume_credits(sql_session, user_id, "essay")
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except Exception:
        raise HTTPException(status_code=500, detail="python-docx is not available in backend environment")

    try:
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                {
                    "role": "user",
                    "parts": [
                        {
                            "text": _essay_prompt(
                                prompt=prompt,
                                essay_format=essay_format,
                                school_name=school_name,
                                course_code=course_code,
                                course_name=course_name,
                                student_name=student_name,
                                instructor_name=instructor_name,
                                due_date=due_date,
                                custom_instructions=custom_instructions,
                            )
                        }
                    ],
                }
            ],
            config={"response_mime_type": "text/plain"},
        )

        generated_text = (getattr(response, "text", None) or "").strip()
        if not generated_text:
            raise HTTPException(status_code=502, detail="Model returned empty essay content")

        title, all_paragraphs = _split_title_and_body(generated_text)
        body_paragraphs, references_paragraphs = _split_body_and_references(all_paragraphs)

        doc: "Document"
        if essay_format == "CUSTOM" and template_docx is not None:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_template:
                tmp_template.write(await template_docx.read())
                tmp_path = tmp_template.name
            try:
                doc = Document(tmp_path)
            finally:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
        else:
            doc = Document()

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        if essay_format == "APA":
            cover_lines = [
                title,
                student_name or "",
                school_name or "",
                " ".join([v for v in [course_code, course_name] if v]).strip(),
                instructor_name or "",
                due_date or "",
            ]
            for idx, line in enumerate(cover_lines):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if idx == 0:
                    p.runs[0].bold = True
            doc.add_page_break()
            heading = doc.add_paragraph(title)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if heading.runs:
                heading.runs[0].bold = True
        elif essay_format == "MLA":
            mla_header = [
                student_name or "",
                instructor_name or "",
                " ".join([v for v in [course_code, course_name] if v]).strip(),
                due_date or "",
            ]
            for line in mla_header:
                if line:
                    doc.add_paragraph(line)
            heading = doc.add_paragraph(title)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if heading.runs:
                heading.runs[0].bold = True
        else:
            # CUSTOM format appends generated content at the end using the template's styles.
            if template_docx is not None:
                doc.add_page_break()
            heading = doc.add_paragraph(title)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if heading.runs:
                heading.runs[0].bold = True

        for para in body_paragraphs:
            p = doc.add_paragraph(para)
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 2.0

        if references_paragraphs:
            doc.add_page_break()
            references_heading = "Works Cited" if essay_format == "MLA" else "References"
            ref_title = doc.add_paragraph(references_heading)
            ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ref_title.runs:
                ref_title.runs[0].bold = True

            for ref in references_paragraphs:
                ref_para = doc.add_paragraph(ref)
                ref_para.paragraph_format.first_line_indent = Inches(-0.5)
                ref_para.paragraph_format.left_indent = Inches(0.5)
                ref_para.paragraph_format.line_spacing = 2.0

        filename_base = "".join(ch for ch in title if ch.isalnum() or ch in (" ", "-", "_")).strip()
        if not filename_base:
            filename_base = "essay"
        filename = f"{filename_base[:100]}.docx"

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Essay generation failed: {str(e)}")
